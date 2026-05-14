import os
import subprocess
import tempfile
import requests
from typing import Optional
import pdfplumber
from docx import Document
from io import BytesIO
import base64

LEVER_API_KEY = os.environ.get("LEVER_API_KEY", "")
LEVER_API_BASE = "https://api.lever.co/v1"


def get_auth_header() -> dict:
    """Get the authorization header for Lever API requests."""
    credentials = base64.b64encode(f"{LEVER_API_KEY}:".encode()).decode()
    return {
        "Authorization": f"Basic {credentials}",
        "Content-Type": "application/json"
    }


def fetch_all_postings() -> list[dict]:
    """Fetch all open and closed postings from Lever."""
    postings = []
    offset = None
    
    while True:
        params = {"limit": 100, "mode": "all"}
        if offset:
            params["offset"] = offset
            
        response = requests.get(
            f"{LEVER_API_BASE}/postings",
            headers=get_auth_header(),
            params=params
        )
        
        if response.status_code != 200:
            raise Exception(f"Failed to fetch postings: {response.text}")
        
        data = response.json()
        postings.extend(data.get("data", []))
        
        if not data.get("hasNext"):
            break
        offset = data.get("next")
    
    return postings


def fetch_candidates_for_posting(posting_id: str, include_active: bool = True, include_archived: bool = False) -> list[dict]:
    """
    Fetch candidates for a specific posting.

    Args:
        posting_id: The Lever posting ID
        include_active: If True, fetch active (non-archived) candidates.
        include_archived: If True, fetch archived candidates.

    Returns:
        List of candidate dictionaries
    """
    candidates = []

    if include_active:
        candidates.extend(_fetch_candidates_with_status(posting_id, archived=False))

    if include_archived:
        candidates.extend(_fetch_candidates_with_status(posting_id, archived=True))

    return candidates


def fetch_all_stages() -> list[dict]:
    """Fetch all stages defined in the Lever account."""
    stages = []
    offset = None

    while True:
        params = {"limit": 100}
        if offset:
            params["offset"] = offset

        response = requests.get(
            f"{LEVER_API_BASE}/stages",
            headers=get_auth_header(),
            params=params
        )

        if response.status_code != 200:
            raise Exception(f"Failed to fetch stages: {response.text}")

        data = response.json()
        stages.extend(data.get("data", []))

        if not data.get("hasNext"):
            break
        offset = data.get("next")

    return stages


def fetch_archive_reasons() -> list[dict]:
    """Fetch all archive reasons defined in the Lever account."""
    reasons = []
    offset = None

    while True:
        params = {"limit": 100}
        if offset:
            params["offset"] = offset

        response = requests.get(
            f"{LEVER_API_BASE}/archive_reasons",
            headers=get_auth_header(),
            params=params,
        )
        if response.status_code != 200:
            raise Exception(f"Failed to fetch archive reasons: {response.text}")

        data = response.json()
        reasons.extend(data.get("data", []))

        if not data.get("hasNext"):
            break
        offset = data.get("next")

    return reasons


def change_candidate_stage(opportunity_id: str, stage_id: str, perform_as: str) -> dict:
    """Move an opportunity to a new stage. Returns the updated opportunity."""
    response = requests.post(
        f"{LEVER_API_BASE}/opportunities/{opportunity_id}/stage",
        headers=get_auth_header(),
        params={"perform_as": perform_as},
        json={"stage": stage_id},
    )
    if response.status_code not in (200, 201):
        raise Exception(f"Failed to change stage (HTTP {response.status_code}): {response.text}")
    return response.json().get("data", {})


def archive_candidate(opportunity_id: str, reason_id: str, perform_as: str) -> dict:
    """Archive an opportunity with the given reason. Returns the updated opportunity."""
    response = requests.post(
        f"{LEVER_API_BASE}/opportunities/{opportunity_id}/archived",
        headers=get_auth_header(),
        params={"perform_as": perform_as},
        json={"reason": reason_id},
    )
    if response.status_code not in (200, 201):
        raise Exception(f"Failed to archive candidate (HTTP {response.status_code}): {response.text}")
    return response.json().get("data", {})


def _fetch_candidates_with_status(posting_id: str, archived: bool) -> list[dict]:
    """
    Internal function to fetch candidates with a specific archived status.

    Args:
        posting_id: The Lever posting ID
        archived: True to fetch archived candidates, False to fetch active candidates

    Returns:
        List of candidate dictionaries
    """
    candidates = []
    offset = None

    while True:
        params = {
            "posting_id": posting_id,
            "limit": 100,
            "archived": "true" if archived else "false"
        }

        if offset:
            params["offset"] = offset

        response = requests.get(
            f"{LEVER_API_BASE}/opportunities",
            headers=get_auth_header(),
            params=params
        )

        if response.status_code != 200:
            raise Exception(f"Failed to fetch candidates: {response.text}")

        data = response.json()
        candidates.extend(data.get("data", []))

        if not data.get("hasNext"):
            break
        offset = data.get("next")

    return candidates


def fetch_candidate_resumes(opportunity_id: str) -> list[dict]:
    """Fetch resume files for a candidate opportunity."""
    response = requests.get(
        f"{LEVER_API_BASE}/opportunities/{opportunity_id}/resumes",
        headers=get_auth_header()
    )
    
    if response.status_code != 200:
        return []
    
    return response.json().get("data", [])


def _detect_format(content: bytes, content_type: str, url_lower: str) -> Optional[str]:
    """Identify resume file format from url, content-type header, and magic bytes."""
    # 1) URL extension (most reliable when Lever exposes it)
    ext_map = {
        ".pdf": "pdf", ".docx": "docx", ".doc": "doc",
        ".rtf": "rtf", ".odt": "odt",
        ".html": "html", ".htm": "html",
        ".txt": "txt", ".text": "txt",
        ".jpg": "image", ".jpeg": "image", ".png": "image",
        ".gif": "image", ".bmp": "image", ".tiff": "image",
        ".tif": "image", ".webp": "image",
    }
    for ext, fmt in ext_map.items():
        if url_lower.endswith(ext):
            return fmt

    # 2) Content-Type header
    if "pdf" in content_type:
        return "pdf"
    if "wordprocessingml" in content_type:
        return "docx"
    if "msword" in content_type or "vnd.ms-word" in content_type:
        return "doc"
    if "rtf" in content_type:
        return "rtf"
    if "opendocument.text" in content_type:
        return "odt"
    if "html" in content_type:
        return "html"
    if content_type.startswith("image/"):
        return "image"
    if content_type.startswith("text/"):
        return "txt"

    # 3) Magic bytes
    if content[:5] == b"%PDF-":
        return "pdf"
    if content[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":  # OLE2 — .doc/.xls/.ppt
        return "doc"
    if content[:5] == b"{\\rtf":
        return "rtf"
    if content[:3] == b"\xff\xd8\xff":  # JPEG
        return "image"
    if content[:8] == b"\x89PNG\r\n\x1a\n":  # PNG
        return "image"
    if content[:6] in (b"GIF87a", b"GIF89a"):
        return "image"
    if content[:4] == b"PK\x03\x04":
        # Office Open XML / ODT / other ZIP. Peek inside for an ODT signature.
        if b"opendocument.text" in content[:1000]:
            return "odt"
        return "docx"  # default ZIP guess — DOCX is overwhelmingly the most common
    sniff = content[:200].lower().lstrip()
    if sniff.startswith(b"<!doctype html") or sniff.startswith(b"<html"):
        return "html"

    return None


def _parse_pdf(content: bytes) -> Optional[str]:
    """Extract text from a PDF using pdfplumber. Returns None for image-only PDFs."""
    try:
        with pdfplumber.open(BytesIO(content)) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text.strip() or None
    except Exception as e:
        print(f"Error parsing PDF: {e}")
        return None


def _parse_docx(content: bytes) -> Optional[str]:
    try:
        doc = Document(BytesIO(content))
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts).strip() or None
    except Exception as e:
        print(f"Error parsing DOCX: {e}")
        return None


def _parse_doc(content: bytes) -> Optional[str]:
    """Parse legacy .doc binary Word format via antiword (system tool)."""
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as f:
            f.write(content)
            tmp_path = f.name
        result = subprocess.run(
            ["antiword", tmp_path],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0:
            return result.stdout.strip() or None
        print(f"antiword failed (rc={result.returncode}): {result.stderr.strip()}")
        return None
    except FileNotFoundError:
        print("antiword is not installed; .doc files cannot be parsed.")
        return None
    except Exception as e:
        print(f"Error parsing DOC: {e}")
        return None
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


def _parse_rtf(content: bytes) -> Optional[str]:
    try:
        from striprtf.striprtf import rtf_to_text
        text = rtf_to_text(content.decode("utf-8", errors="replace"))
        return text.strip() or None
    except Exception as e:
        print(f"Error parsing RTF: {e}")
        return None


def _parse_odt(content: bytes) -> Optional[str]:
    try:
        from odf.opendocument import load
        from odf import teletype
        doc = load(BytesIO(content))
        text = teletype.extractText(doc)
        return text.strip() or None
    except Exception as e:
        print(f"Error parsing ODT: {e}")
        return None


def _parse_html(content: bytes) -> Optional[str]:
    try:
        from bs4 import BeautifulSoup
        try:
            soup = BeautifulSoup(content, "lxml")
        except Exception:
            soup = BeautifulSoup(content, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return "\n".join(lines) or None
    except Exception as e:
        print(f"Error parsing HTML: {e}")
        return None


def _vision_extract_image(image_bytes: bytes, mime_type: str = "image/png") -> Optional[str]:
    """Use GPT-4o vision to OCR a single image."""
    try:
        from openai import OpenAI
        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
        b64 = base64.b64encode(image_bytes).decode("utf-8")
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": "This is a page from a candidate's resume. Extract ALL visible text accurately, preserving section order (contact info, experience, education, skills) and bullet structure. Return only the extracted text — no commentary or markdown wrappers."},
                    {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{b64}"}},
                ],
            }],
            max_tokens=4096,
        )
        text = (response.choices[0].message.content or "").strip()
        return text or None
    except Exception as e:
        print(f"Vision extraction failed: {e}")
        return None


def _vision_extract_pdf(content: bytes) -> Optional[str]:
    """Render each PDF page to PNG and OCR with vision. Used when pdfplumber returns no text."""
    try:
        import fitz  # PyMuPDF
        pages_text = []
        with fitz.open(stream=content, filetype="pdf") as pdf:
            for page in pdf:
                pix = page.get_pixmap(dpi=150)
                page_text = _vision_extract_image(pix.tobytes("png"), "image/png")
                if page_text:
                    pages_text.append(page_text)
        return "\n\n".join(pages_text).strip() or None
    except Exception as e:
        print(f"Vision PDF extraction failed: {e}")
        return None


def download_and_parse_resume(file_download_url: str) -> Optional[str]:
    """Download a resume file and extract text content.

    Handles PDF (text + scanned via vision), DOCX, DOC (via antiword),
    RTF, ODT, HTML, plain text, and images (via vision). Returns None for
    truly unparseable formats so the caller skips the candidate rather
    than feeding garbage to the LLM.
    """
    try:
        response = requests.get(file_download_url, headers=get_auth_header())
        if response.status_code != 200:
            return None

        content = response.content
        content_type = response.headers.get("Content-Type", "").lower()
        url_lower = file_download_url.lower()
        fmt = _detect_format(content, content_type, url_lower)

        if fmt == "pdf":
            text = _parse_pdf(content)
            if text:
                return text
            # No extractable text — likely a scanned PDF. Fall back to vision.
            print("PDF had no extractable text; falling back to vision OCR.")
            return _vision_extract_pdf(content)

        if fmt == "docx":
            return _parse_docx(content)

        if fmt == "doc":
            return _parse_doc(content)

        if fmt == "rtf":
            return _parse_rtf(content)

        if fmt == "odt":
            return _parse_odt(content)

        if fmt == "html":
            return _parse_html(content)

        if fmt == "txt":
            return response.text.strip() or None

        if fmt == "image":
            mime = content_type if content_type.startswith("image/") else "image/png"
            return _vision_extract_image(content, mime)

        # Unknown — last-ditch: if it looks like plain text, decode it.
        try:
            decoded = content.decode("utf-8")
            if decoded.isprintable() or "\n" in decoded:
                return decoded.strip() or None
        except UnicodeDecodeError:
            pass

        print(f"Unrecognized resume format (content-type={content_type!r}, url={file_download_url})")
        return None

    except Exception as e:
        print(f"Error parsing resume: {e}")
        return None


def get_candidate_linkedin(candidate: dict) -> Optional[str]:
    """Extract LinkedIn URL from candidate data."""
    links = candidate.get("links", [])
    for link in links:
        if "linkedin" in link.lower():
            return link
    
    urls = candidate.get("urls", {})
    if isinstance(urls, dict):
        linkedin = urls.get("linkedin")
        if linkedin:
            return linkedin
    
    return None


def get_candidate_lever_url(candidate: dict) -> str:
    """Generate the Lever profile URL for a candidate."""
    opportunity_id = candidate.get("id", "")
    return f"https://hire.lever.co/candidates/{opportunity_id}"


def get_candidate_name(candidate: dict) -> str:
    """Get the candidate's name."""
    return candidate.get("name", "Unknown Candidate")


def get_candidate_email(candidate: dict) -> Optional[str]:
    """Get the candidate's email."""
    emails = candidate.get("emails", [])
    return emails[0] if emails else None


def get_resume_text_for_candidate(opportunity_id: str) -> Optional[str]:
    """Get parsed resume text for a candidate."""
    resumes = fetch_candidate_resumes(opportunity_id)
    
    for resume in resumes:
        file_info = resume.get("file") or {}
        download_url = file_info.get("downloadUrl")
        
        if download_url:
            text = download_and_parse_resume(download_url)
            if text:
                return text
    
    return None
